"""
Export Service for IEEE Report Restructurer.
Handles generation of IEEE-formatted DOCX and PDF documents.
"""

import os
import re
import logging
from typing import List, Optional
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from ..models.section import Section, IEEECategory
from ..models.document import Document, GlobalContext
from ..config import get_settings

# Set up logging
logger = logging.getLogger(__name__)


class ExportService:
    """Service for exporting documents to DOCX and PDF formats."""
    
    def __init__(self):
        self.settings = get_settings()
        self._ensure_output_dir()
    
    def _ensure_output_dir(self):
        """Ensure output directory exists."""
        os.makedirs(self.settings.output_dir, exist_ok=True)
    
    def generate_docx(
        self,
        document: Document,
        sections: List[Section],
        context: GlobalContext,
    ) -> str:
        """
        Generate IEEE-formatted DOCX document.
        This is the PRIMARY guaranteed output.
        
        Args:
            document: Document model
            sections: Processed sections
            context: Global context
            
        Returns:
            Path to generated DOCX file
        """
        doc = DocxDocument()
        
        # Set up IEEE styles
        self._setup_ieee_styles(doc)
        
        # Set page margins (IEEE typically uses narrow margins)
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.65)
            section.right_margin = Inches(0.65)
        
        # Add title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(context.project_title)
        title_run.bold = True
        title_run.font.size = Pt(24)
        title_run.font.name = 'Times New Roman'
        
        # Add authors if available
        if context.authors:
            authors_para = doc.add_paragraph()
            authors_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            authors_run = authors_para.add_run(", ".join(context.authors))
            authors_run.font.size = Pt(11)
            authors_run.font.name = 'Times New Roman'
            authors_run.italic = True
        
        # Add a line break
        doc.add_paragraph()
        
        # Process each section
        for section in sections:
            self._add_section_to_doc(doc, section)
        
        # Save document
        output_path = os.path.join(
            self.settings.output_dir,
            f"{document.id}_ieee.docx"
        )
        doc.save(output_path)
        
        logger.info(f"DOCX generated successfully: {output_path}")
        return output_path
    
    def _setup_ieee_styles(self, doc: DocxDocument):
        """Set up IEEE-compatible styles in the document."""
        styles = doc.styles
        
        # Main heading style (I., II., etc.)
        try:
            heading1_style = styles.add_style('IEEE Heading 1', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            heading1_style = styles['IEEE Heading 1']
        
        heading1_style.font.name = 'Times New Roman'
        heading1_style.font.size = Pt(10)
        heading1_style.font.bold = True
        heading1_style.font.all_caps = True
        heading1_style.paragraph_format.space_before = Pt(12)
        heading1_style.paragraph_format.space_after = Pt(6)
        heading1_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subheading style (A., B., etc.)
        try:
            heading2_style = styles.add_style('IEEE Heading 2', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            heading2_style = styles['IEEE Heading 2']
        
        heading2_style.font.name = 'Times New Roman'
        heading2_style.font.size = Pt(10)
        heading2_style.font.italic = True
        heading2_style.paragraph_format.space_before = Pt(10)
        heading2_style.paragraph_format.space_after = Pt(4)
        
        # Body text style
        try:
            body_style = styles.add_style('IEEE Body', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            body_style = styles['IEEE Body']
        
        body_style.font.name = 'Times New Roman'
        body_style.font.size = Pt(10)
        body_style.paragraph_format.first_line_indent = Inches(0.25)
        body_style.paragraph_format.space_after = Pt(0)
        body_style.paragraph_format.line_spacing = 1.0
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Abstract style
        try:
            abstract_style = styles.add_style('IEEE Abstract', WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            abstract_style = styles['IEEE Abstract']
        
        abstract_style.font.name = 'Times New Roman'
        abstract_style.font.size = Pt(9)
        abstract_style.font.bold = True
        abstract_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    def _add_section_to_doc(self, doc: DocxDocument, section: Section):
        """Add a section to the document."""
        cat = IEEECategory(section.category) if isinstance(section.category, str) else section.category
        content = section.rewritten_content or section.original_content
        
        # Handle special sections
        if cat == IEEECategory.ABSTRACT:
            # Abstract heading
            heading_para = doc.add_paragraph()
            heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            heading_run = heading_para.add_run("Abstract")
            heading_run.bold = True
            heading_run.italic = True
            heading_run.font.size = Pt(9)
            heading_run.font.name = 'Times New Roman'
            
            # Abstract content
            abs_para = doc.add_paragraph()
            abs_run = abs_para.add_run(content)
            abs_run.font.size = Pt(9)
            abs_run.font.name = 'Times New Roman'
            abs_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
        elif cat == IEEECategory.KEYWORDS:
            # Keywords
            kw_para = doc.add_paragraph()
            kw_run = kw_para.add_run(content)
            kw_run.font.size = Pt(9)
            kw_run.font.name = 'Times New Roman'
            kw_run.italic = True
            
        elif cat == IEEECategory.REFERENCES:
            # References heading
            ref_heading = doc.add_paragraph()
            ref_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ref_run = ref_heading.add_run("REFERENCES")
            ref_run.bold = True
            ref_run.font.size = Pt(10)
            ref_run.font.name = 'Times New Roman'
            
            # Each reference on its own line
            for line in content.split('\n'):
                if line.strip():
                    ref_para = doc.add_paragraph()
                    ref_para.add_run(line.strip()).font.size = Pt(8)
                    ref_para.paragraph_format.left_indent = Inches(0.25)
                    ref_para.paragraph_format.first_line_indent = Inches(-0.25)
                    
        else:
            # Regular section
            # Add heading
            heading_text = self._format_heading(section)
            
            if section.level == 1:
                heading_para = doc.add_paragraph()
                heading_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                heading_run = heading_para.add_run(heading_text.upper())
                heading_run.bold = True
                heading_run.font.size = Pt(10)
                heading_run.font.name = 'Times New Roman'
            else:
                heading_para = doc.add_paragraph()
                heading_run = heading_para.add_run(heading_text)
                heading_run.italic = True
                heading_run.font.size = Pt(10)
                heading_run.font.name = 'Times New Roman'
            
            # Add content paragraphs
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    body_para = doc.add_paragraph()
                    body_run = body_para.add_run(para_text.strip())
                    body_run.font.size = Pt(10)
                    body_run.font.name = 'Times New Roman'
                    body_para.paragraph_format.first_line_indent = Inches(0.25)
                    body_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    body_para.paragraph_format.line_spacing = 1.0
    
    def _format_heading(self, section: Section) -> str:
        """Format section heading with number."""
        number = section.ieee_number or ""
        title = section.title
        
        if number:
            return f"{number} {title}"
        return title
    
    def generate_pdf_native(
        self,
        document: Document,
        sections: List[Section],
        context: GlobalContext,
    ) -> Optional[str]:
        """
        Generate PDF directly using fpdf2 (pure Python, no external dependencies).
        
        Args:
            document: Document model
            sections: Processed sections
            context: Global context
            
        Returns:
            Path to generated PDF file, or None if generation failed
        """
        try:
            from fpdf import FPDF
            
            # Create PDF
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            
            # Set margins
            pdf.set_margins(left=18, top=20, right=18)
            
            # Clean title for PDF encoding
            title = self._clean_text_for_pdf(context.project_title or "IEEE Document")
            
            # Title
            pdf.set_font("Times", "B", 24)
            pdf.multi_cell(0, 10, title, align="C")
            pdf.ln(5)
            
            # Authors (cleaned)
            if context.authors:
                authors_text = ", ".join([self._clean_text_for_pdf(a) for a in context.authors])
                pdf.set_font("Times", "I", 11)
                pdf.multi_cell(0, 6, authors_text, align="C")
                pdf.ln(5)
            
            pdf.ln(10)
            
            # Process sections
            for section in sections:
                try:
                    self._add_section_to_pdf(pdf, section)
                except Exception as section_error:
                    logger.warning(f"Error adding section '{section.title}' to PDF: {section_error}")
                    # Continue with other sections
            
            # Ensure output directory exists
            os.makedirs(self.settings.output_dir, exist_ok=True)
            
            # Save PDF
            output_path = os.path.join(
                self.settings.output_dir,
                f"{document.id}_ieee.pdf"
            )
            pdf.output(output_path)
            
            logger.info(f"PDF generated successfully with fpdf2: {output_path}")
            return output_path
            
        except ImportError as e:
            logger.warning(f"fpdf2 not installed. Cannot generate PDF: {e}")
            return None
        except Exception as e:
            logger.error(f"PDF generation with fpdf2 failed: {e}", exc_info=True)
            return None
    
    def _add_section_to_pdf(self, pdf, section: Section):
        """Add a section to the PDF document."""
        cat = IEEECategory(section.category) if isinstance(section.category, str) else section.category
        content = section.rewritten_content or section.original_content
        
        # Clean content for PDF (remove problematic characters)
        content = self._clean_text_for_pdf(content)
        
        if cat == IEEECategory.ABSTRACT:
            # Abstract heading
            pdf.set_font("Times", "BI", 9)
            pdf.multi_cell(0, 5, "Abstract", align="C")
            pdf.ln(2)
            
            # Abstract content
            pdf.set_font("Times", "", 9)
            pdf.multi_cell(0, 5, content)
            pdf.ln(5)
            
        elif cat == IEEECategory.KEYWORDS:
            pdf.set_font("Times", "I", 9)
            pdf.multi_cell(0, 5, content)
            pdf.ln(5)
            
        elif cat == IEEECategory.REFERENCES:
            # References heading
            pdf.set_font("Times", "B", 10)
            pdf.multi_cell(0, 6, "REFERENCES", align="C")
            pdf.ln(3)
            
            # Each reference
            pdf.set_font("Times", "", 8)
            for line in content.split('\n'):
                if line.strip():
                    pdf.multi_cell(0, 4, line.strip())
                    pdf.ln(1)
            pdf.ln(3)
            
        else:
            # Regular section heading
            heading_text = self._format_heading(section)
            
            if section.level == 1:
                pdf.set_font("Times", "B", 10)
                pdf.multi_cell(0, 6, heading_text.upper(), align="C")
            else:
                pdf.set_font("Times", "I", 10)
                pdf.multi_cell(0, 6, heading_text)
            
            pdf.ln(2)
            
            # Content paragraphs
            pdf.set_font("Times", "", 10)
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    # Add indent for first line effect
                    pdf.multi_cell(0, 5, "    " + para_text.strip())
                    pdf.ln(2)
            
            pdf.ln(3)
    
    def _clean_text_for_pdf(self, text: str) -> str:
        """Clean text for PDF rendering (handle encoding issues, HTML, placeholders)."""
        if not text:
            return ""
        
        # Strip HTML artifacts
        html_patterns = [
            r'<hr\s*/?>',
            r'<br\s*/?>',
            r'</?div[^>]*>',
            r'</?span[^>]*>',
            r'</?p[^>]*>',
            r'</?strong>',
            r'</?em>',
            r'</?b>',
            r'</?i>',
            r'</?u>',
            r'&nbsp;',
            r'&amp;',
            r'&lt;',
            r'&gt;',
            r'&quot;',
        ]
        
        for pattern in html_patterns:
            text = re.sub(pattern, ' ', text, flags=re.IGNORECASE)
        
        # Replace placeholders
        placeholder_patterns = [
            (r'\[Your Name\]', ''),
            (r'\[Author Name\]', ''),
            (r'\[Name\]', ''),
            (r'\[Date\]', ''),
            (r'\[Insert .+?\]', ''),
            (r'\[TODO.*?\]', ''),
            (r'\[TBD.*?\]', ''),
            (r'\[PLACEHOLDER.*?\]', ''),
        ]
        
        for pattern, replacement in placeholder_patterns:
            text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
        
        # Replace problematic unicode characters
        replacements = {
            '\u2018': "'",  # Left single quote
            '\u2019': "'",  # Right single quote
            '\u201c': '"',  # Left double quote
            '\u201d': '"',  # Right double quote
            '\u2013': '-',  # En dash
            '\u2014': '--', # Em dash
            '\u2026': '...', # Ellipsis
            '\u00a0': ' ',  # Non-breaking space
            '\u2022': '*',  # Bullet
            '\u00b7': '*',  # Middle dot
            '\u2212': '-',  # Minus sign
            '\u00d7': 'x',  # Multiplication sign
            '\u00f7': '/',  # Division sign
            '\u2192': '->',  # Right arrow
            '\u2190': '<-',  # Left arrow
            '\u2713': '[x]', # Check mark
            '\u2717': '[ ]', # X mark
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        # Clean up multiple spaces/newlines
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove any remaining non-ASCII characters that might cause issues
        try:
            text.encode('latin-1')
        except UnicodeEncodeError:
            text = text.encode('latin-1', errors='replace').decode('latin-1')
        
        return text.strip()
    
    def generate_pdf(self, docx_path: str) -> Optional[str]:
        """
        Convert DOCX to PDF using docx2pdf (requires MS Word/LibreOffice).
        This is a fallback method.
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            Path to generated PDF file, or None if conversion failed
        """
        if not docx_path or not os.path.exists(docx_path):
            logger.warning(f"DOCX file not found for PDF conversion: {docx_path}")
            return None
        
        pdf_path = docx_path.replace('.docx', '.pdf')
        
        # Method 1: Try using docx2pdf (requires Microsoft Word on Windows)
        try:
            from docx2pdf import convert
            logger.info(f"Attempting PDF conversion with docx2pdf: {docx_path}")
            convert(docx_path, pdf_path)
            
            if os.path.exists(pdf_path):
                logger.info(f"PDF generated successfully with docx2pdf: {pdf_path}")
                return pdf_path
                
        except ImportError:
            logger.warning("docx2pdf not installed.")
        except Exception as e:
            logger.warning(f"docx2pdf conversion failed: {e}")
        
        # Method 2: Try using LibreOffice as fallback
        try:
            import subprocess
            result = subprocess.run(
                ['soffice', '--headless', '--convert-to', 'pdf',
                 '--outdir', os.path.dirname(docx_path), docx_path],
                capture_output=True, timeout=60
            )
            
            if os.path.exists(pdf_path):
                logger.info(f"PDF generated successfully with LibreOffice: {pdf_path}")
                return pdf_path
                
        except Exception as e:
            logger.warning(f"LibreOffice conversion failed: {e}")
        
        return None
    
    def generate_both_formats(
        self,
        document: Document,
        sections: List[Section],
        context: GlobalContext,
    ) -> dict:
        """
        Generate both DOCX and PDF formats.
        DOCX is guaranteed; PDF uses native generation (fpdf2).
        
        Args:
            document: Document model
            sections: Processed sections
            context: Global context
            
        Returns:
            Dict with paths to generated files
        """
        result = {
            "docx_path": None,
            "pdf_path": None,
            "errors": [],
        }
        
        # Step 1: Generate DOCX (primary, required output)
        try:
            result["docx_path"] = self.generate_docx(document, sections, context)
            logger.info(f"DOCX export successful: {result['docx_path']}")
        except Exception as e:
            error_msg = f"DOCX generation failed: {str(e)}"
            logger.error(error_msg)
            result["errors"].append(error_msg)
            return result
        
        # Step 2: Generate PDF using native fpdf2 (no external dependencies)
        try:
            pdf_path = self.generate_pdf_native(document, sections, context)
            if pdf_path and os.path.exists(pdf_path):
                result["pdf_path"] = pdf_path
                logger.info(f"PDF export successful: {pdf_path}")
            else:
                # Fallback: try docx2pdf conversion
                pdf_path = self.generate_pdf(result["docx_path"])
                if pdf_path and os.path.exists(pdf_path):
                    result["pdf_path"] = pdf_path
                else:
                    result["errors"].append("PDF generation not available.")
        except Exception as e:
            logger.warning(f"PDF generation failed: {e}")
            result["errors"].append(f"PDF generation failed: {str(e)}")
        
        return result
